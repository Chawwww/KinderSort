"""
docx_export.py — Convert guidebook.md to KinderSort_Teacher_Guide.docx.

Parses guidebook.md and builds a styled Word document with inline screenshots.
Screenshots must exist in guidebook_assets/ before running this script.

Usage:
    pip install python-docx
    python docx_export.py
"""

from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    raise SystemExit(
        "python-docx is not installed.\nRun: pip install python-docx"
    )


def build_docx(md_path: Path, out_path: Path) -> None:
    """Convert a markdown file to a .docx document.

    Handles: headings (#/##/###/####), paragraphs, bullet lists, numbered lists,
    fenced code blocks, inline images (![alt](path)), horizontal rules, blockquotes.

    Args:
        md_path: Path to the input guidebook.md file.
        out_path: Path where the .docx will be written.
    """
    doc = Document()

    # --- Document-level styling ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Tighten default heading spacing
    for level in range(1, 5):
        try:
            h_style = doc.styles[f"Heading {level}"]
            h_style.font.name = "Calibri"
        except KeyError:
            pass

    content = md_path.read_text(encoding="utf-8")
    lines = content.split("\n")

    in_code_block = False
    code_lines: list[str] = []

    for raw_line in lines:
        line = raw_line.rstrip()

        # ---- Fenced code block ----
        if line.startswith("```"):
            if in_code_block:
                # Flush collected code lines
                if code_lines:
                    para = doc.add_paragraph()
                    run = para.add_run("\n".join(code_lines))
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    para.paragraph_format.left_indent = Inches(0.5)
                    para.paragraph_format.space_before = Pt(4)
                    para.paragraph_format.space_after = Pt(4)
                code_lines = []
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # ---- Headings ----
        if line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)

        # ---- Inline image  ![alt](path) ----
        elif line.startswith("!["):
            try:
                # Extract path between ( and )
                img_path_str = line.split("(", 1)[1].rstrip(")")
                img_path = Path(img_path_str)
                if img_path.exists():
                    para = doc.add_paragraph()
                    run = para.add_run()
                    run.add_picture(str(img_path), width=Inches(5.5))
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.space_before = Pt(6)
                    para.paragraph_format.space_after = Pt(6)
                else:
                    doc.add_paragraph(f"[Screenshot not yet available: {img_path_str}]")
            except Exception as exc:  # noqa: BLE001
                doc.add_paragraph(f"[Image error: {exc}]")

        # ---- Horizontal rule ----
        elif line.strip() == "---":
            para = doc.add_paragraph()
            pf = para.paragraph_format
            pf.space_before = Pt(4)
            pf.space_after = Pt(4)
            # Add a visual separator via a border (top border on an empty paragraph)
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)

        # ---- Bullet list ----
        elif line.startswith("  - ") or line.startswith("  * "):
            doc.add_paragraph(line[4:], style="List Bullet 2")
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")

        # ---- Numbered list ----
        elif len(line) > 2 and line[0].isdigit() and line[1] in ".)" and line[2] == " ":
            doc.add_paragraph(line[3:], style="List Number")

        # ---- Blockquote ----
        elif line.startswith("> "):
            para = doc.add_paragraph(line[2:])
            para.paragraph_format.left_indent = Inches(0.4)
            for run in para.runs:
                run.italic = True

        # ---- Empty line ----
        elif not line.strip():
            pass  # Word adds natural spacing between paragraphs

        # ---- Normal paragraph ----
        else:
            # Strip leading * for italic/bold markers (simplified)
            doc.add_paragraph(line)

    doc.save(str(out_path))
    print(f"Saved: {out_path.resolve()}")
    print(f"  Size: {out_path.stat().st_size:,} bytes")


if __name__ == "__main__":
    md_path = Path("guidebook.md")
    out_path = Path("KinderSort_Teacher_Guide.docx")

    if not md_path.exists():
        raise SystemExit(f"guidebook.md not found at {md_path.resolve()}")

    print("Converting {} -> {}...".format(md_path, out_path))
    build_docx(md_path, out_path)
    print("Done!")
