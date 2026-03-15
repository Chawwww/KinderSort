"""
generate_guide.py — Automated screenshot guidebook generator for KinderSort.

Launches the KinderSort GUI as a subprocess, automates UI interactions using
pywinauto, captures screenshots at key states, then writes guidebook.md with
embedded image references and an optional .docx export.

Requirements:
    pip install pyautogui pywinauto pillow

Usage:
    python generate_guide.py
"""

import subprocess
import sys
import time
from pathlib import Path

import pyautogui
import pywinauto
from pywinauto import Application

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

ASSETS_DIR = Path("guidebook_assets")
PYTHON_EXE = Path(".venv/Scripts/python.exe")
APP_TITLE = "KinderSort — Student Photo Organiser"
REF_FOLDER = str(Path("referencePhoto").resolve())
EVENTS_FOLDER = str(Path("Events").resolve())
OUTPUT_FOLDER = str(Path("Output").resolve())

# Paths to pre-run screenshots (if app already ran)
WAIT_SECONDS = 4  # Time to wait for window to fully render


def setup_assets_dir() -> None:
    """Create the guidebook_assets directory if it doesn't exist."""
    ASSETS_DIR.mkdir(exist_ok=True)
    print(f"Assets directory: {ASSETS_DIR.resolve()}")


def wait_for_window(title: str, timeout: int = 20) -> Application:
    """Poll until a window with the given title appears, then return it.

    Uses win32 backend which works reliably with tkinter applications.

    Args:
        title: Exact window title to search for.
        timeout: Seconds to wait before raising TimeoutError.

    Returns:
        Connected pywinauto Application instance.

    Raises:
        TimeoutError: If window doesn't appear within timeout seconds.
    """
    print(f"Waiting for window '{title}'...")
    deadline = time.time() + timeout
    while time.time() < deadline:
        try:
            app = Application(backend="win32").connect(title=title, timeout=2)
            print("Window found!")
            return app
        except Exception:  # noqa: BLE001
            time.sleep(1)
    raise TimeoutError(f"Window '{title}' did not appear within {timeout}s")


def screenshot(name: str) -> Path:
    """Take a screenshot and save to guidebook_assets/{name}.

    Args:
        name: Filename without extension (e.g. '01_launch').

    Returns:
        Path to the saved screenshot file.
    """
    time.sleep(0.5)  # Brief pause to let UI settle
    path = ASSETS_DIR / f"{name}.png"
    pyautogui.screenshot(str(path))
    print(f"  Screenshot saved: {path.name}")
    return path


def fill_folder_entry(win, row_index: int, folder_path: str) -> None:
    """Click the Browse button for a folder row and handle the file dialog.

    In tkinter, all button labels are empty strings via win32 API (tkinter doesn't
    expose button text via Win32 WM_GETTEXT). Buttons are in creation order:
    index 0=Browse ref, 1=Browse events, 2=Browse output, 3=Start, 4=Cancel.

    Args:
        win: pywinauto window wrapper (win32 backend).
        row_index: 0-based index of the folder row (0=Reference, 1=Events, 2=Output).
        folder_path: Absolute path string to enter in the dialog.
    """
    try:
        all_buttons = win.descendants(class_name="Button")
        # Button order in tkinter: Browse(0), Browse(1), Browse(2), Start(3), Cancel(4)
        if row_index < len(all_buttons):
            all_buttons[row_index].click_input()
            time.sleep(1.5)  # Wait for dialog to open
        else:
            print(f"  Warning: Button index {row_index} out of range ({len(all_buttons)} buttons found)")
            return
    except Exception as e:  # noqa: BLE001
        print(f"  Warning: Button click failed: {e}")
        return

    # Navigate to path in the open tkinter askdirectory dialog
    # tkinter uses the Windows SHELL folder browser — type path into address bar
    try:
        pyautogui.hotkey("alt", "d")  # Focus address bar
        time.sleep(0.4)
        pyautogui.hotkey("ctrl", "a")
        time.sleep(0.1)
        pyautogui.typewrite(folder_path, interval=0.02)
        pyautogui.press("enter")
        time.sleep(1.2)
        pyautogui.press("enter")  # Confirm selection ("Select Folder" / OK)
        time.sleep(0.8)
    except Exception as e:  # noqa: BLE001
        print(f"  Warning: Dialog navigation failed: {e}")


def run_guide_capture() -> None:
    """Main automation routine — launch app, capture screenshots, write guide."""
    setup_assets_dir()

    # -----------------------------------------------------------------------
    # Launch the app
    # -----------------------------------------------------------------------
    print("\n[1/7] Launching KinderSort...")
    proc = subprocess.Popen(
        [str(PYTHON_EXE), "main.py"],
        creationflags=subprocess.CREATE_NEW_PROCESS_GROUP if sys.platform == "win32" else 0,
    )

    try:
        app = wait_for_window(APP_TITLE, timeout=20)
        win = app.window(title=APP_TITLE)
        win.set_focus()
        time.sleep(WAIT_SECONDS)

        # -----------------------------------------------------------------------
        # State 1: Fresh launch
        # -----------------------------------------------------------------------
        print("[1/7] State 1: App on first launch (blank)")
        win.set_focus()
        screenshot("01_launch")

        # -----------------------------------------------------------------------
        # State 2: Set Reference folder
        # -----------------------------------------------------------------------
        print("[2/7] State 2: Selecting Reference Photos folder")
        fill_folder_entry(win, 0, REF_FOLDER)
        win.set_focus()
        screenshot("02_reference_selected")

        # -----------------------------------------------------------------------
        # State 3: Set Events folder
        # -----------------------------------------------------------------------
        print("[3/7] State 3: Selecting Events folder")
        fill_folder_entry(win, 1, EVENTS_FOLDER)
        win.set_focus()
        screenshot("03_events_selected")

        # -----------------------------------------------------------------------
        # State 4: Set Output folder → all three selected (ready state)
        # -----------------------------------------------------------------------
        print("[4/7] State 4: Selecting Output folder (all three selected)")
        fill_folder_entry(win, 2, OUTPUT_FOLDER)
        win.set_focus()
        screenshot("04_all_folders_set")

        # -----------------------------------------------------------------------
        # State 5: Click Start Sorting → capture mid-progress
        # -----------------------------------------------------------------------
        print("[5/7] State 5: Starting sort and capturing progress")
        try:
            all_buttons = win.descendants(class_name="Button")
            # Start Sorting is button index 3 (after 3 Browse buttons)
            if len(all_buttons) >= 4:
                all_buttons[3].click_input()
                print("  Clicked Start Sorting (button index 3)")
            else:
                print(f"  Warning: Not enough buttons found ({len(all_buttons)})")
        except Exception as e:  # noqa: BLE001
            print(f"  Warning: Could not click Start: {e}")

        # Wait ~20s then screenshot mid-progress
        time.sleep(20)
        win.set_focus()
        screenshot("05_sorting_in_progress")

        # -----------------------------------------------------------------------
        # State 6: Wait for completion
        # -----------------------------------------------------------------------
        print("[6/7] State 6: Waiting for sorting to complete...")
        # Poll every 10s for up to 15 minutes
        deadline = time.time() + 900
        while time.time() < deadline:
            time.sleep(10)
            # Check if the Start button (index 3) is re-enabled (indicates completion)
            try:
                all_buttons = win.descendants(class_name="Button")
                if len(all_buttons) >= 4 and all_buttons[3].is_enabled():
                    print("  Sorting complete!")
                    break
            except Exception:  # noqa: BLE001
                pass

        try:
            win.set_focus()
        except Exception:  # noqa: BLE001
            pass
        screenshot("06_sorting_complete")

        print("\n[7/7] All screenshots captured!")
        print(f"Screenshots in: {ASSETS_DIR.resolve()}")

    finally:
        # Don't kill the process — user may want to inspect results
        print(f"\nApp still running (PID {proc.pid}). Close it manually when done.")


def write_guidebook_md() -> None:
    """Write the teacher guidebook as guidebook.md with embedded screenshots."""
    print("\nWriting guidebook.md...")

    content = """# KinderSort — Teacher's Guide

*How to sort your students' event photos automatically*

---

## What KinderSort Does

KinderSort looks at each photo from a school event and finds your students' faces.
It then automatically puts each photo into the right student's folder — so you don't
have to sort hundreds of photos by hand!

**One photo can appear in multiple folders** — for example, a group shot with three
students will be copied to all three students' folders.

---

## Before You Start

You need three folders ready on your computer:

### 1. Reference Photos Folder
One clear, front-facing photo of each student.
- Name each photo with the student's full name
- Examples: `Ali.jpg`, `Siti.png`, `Kumar.jpeg`
- Make sure the face is clearly visible and well-lit

### 2. Events Folder
A folder containing **subfolders** — one subfolder per event.
- Example structure:
  ```
  Events/
      Sports_Day/
          IMG_001.jpg
          IMG_002.jpg
      Field_Trip/
          IMG_003.jpg
  ```
- ⚠️ **Important:** Photos must be inside a subfolder (the event name), not directly in the Events folder.

### 3. Output Folder
An empty folder where KinderSort will save the sorted results.
You can create a new empty folder anywhere on your computer.

---

## Step-by-Step Guide

### Step 1 — Open KinderSort

Double-click the **KinderSort.exe** file to launch the app.

![KinderSort on launch](guidebook_assets/01_launch.png)

You will see three folder selector rows at the top.

---

### Step 2 — Select Your Reference Photos Folder

Click the **Browse…** button next to "Reference Photos".

Navigate to your folder of student reference photos and click **Select Folder**.

![Reference folder selected](guidebook_assets/02_reference_selected.png)

The path to your folder will appear in the box.

---

### Step 3 — Select Your Events Folder

Click the **Browse…** button next to "Events Folder".

Navigate to your Events folder (the one containing event subfolders) and click **Select Folder**.

![Events folder selected](guidebook_assets/03_events_selected.png)

---

### Step 4 — Select Your Output Folder

Click the **Browse…** button next to "Output Folder".

Navigate to your empty output folder and click **Select Folder**.

When all three folders are selected, the app is ready.

![All three folders selected](guidebook_assets/04_all_folders_set.png)

---

### Step 5 — Start Sorting

Click the large green **Start Sorting** button.

KinderSort will begin processing your photos. You will see:
- A **progress bar** filling up as photos are processed
- A **status line** showing the current photo filename

![Sorting in progress](guidebook_assets/05_sorting_in_progress.png)

> ⏱️ Processing time depends on the number of photos. Expect about 1–2 minutes per 10 photos on a normal computer.

You can click **Cancel** at any time to stop — photos processed so far will be saved.

---

### Step 6 — Review the Results

When sorting is complete, you will see a summary:

![Sorting complete](guidebook_assets/06_sorting_complete.png)

The summary shows:
- **Total images found** — how many photos were scanned
- **Matched (sorted)** — how many photos were placed into student folders
- **Unmatched** — photos where no student face was recognised
- **Skipped (errors)** — photos that could not be opened

---

## Understanding Your Output Folder

After sorting, your Output folder will look like this:

```
Output/
    Ali/
        Sports_Day__IMG_001.jpg
        Concert__IMG_045.jpg
    Siti/
        Sports_Day__IMG_001.jpg   ← same group photo, copied here too
        Field_Trip__IMG_023.jpg
    _unmatched/
        blurry_photo.jpg
        background_only.jpg
    kindersort_log.txt
```

**Each student has their own folder** containing all photos where their face was detected.

The **`_unmatched` folder** contains:
- Photos where no faces were detected (e.g. landscape shots, blurry images)
- Photos where faces were detected but didn't match any student (e.g. teachers, parents)

The **`kindersort_log.txt` file** is a detailed record of everything KinderSort did.
You don't need to read it normally, but it's useful if something seems wrong.

---

## Common Problems & Solutions

### "No face was detected" warning on startup
**Cause:** One of your reference photos doesn't show the student's face clearly enough.

**Fix:** Replace that student's reference photo with a clearer, front-facing photo.
Good lighting and a plain background work best.

### Many photos end up in `_unmatched`
**Possible causes:**
1. Reference photo quality is poor — use a clearer photo
2. Event photos are very blurry or taken from far away
3. Students are wearing face masks or hats in the event photos

**Fix:** Try better reference photos first. KinderSort cannot recognise faces that
are partially covered or turned away.

### "Missing folders" error when clicking Start
**Fix:** Make sure all three folder fields are filled in before clicking Start Sorting.

### The app seems stuck / progress bar not moving
**Cause:** Face recognition is CPU-intensive — it is still working.

**Fix:** Wait patiently. A batch of 100 photos may take 10–15 minutes on an older computer.
Watch the status line — if it shows a new filename, it is still running.

### Output folder photos have long names like `Sports_Day__IMG_001.jpg`
This is normal! The event folder name is added as a prefix so you always know
which event a photo came from.

---

## Tips for Better Results

1. **Use a clear, recent reference photo** — front-facing, good lighting, no sunglasses
2. **One face per reference photo** — if a reference photo has multiple faces, KinderSort uses the first face detected
3. **Consistent lighting** helps — very dark or backlit event photos may not match well
4. **Re-run with a higher tolerance** if you're missing matches — ask your IT support to adjust the `DISTANCE_THRESHOLD` in `sorter.py` from `0.5` to `0.6`
5. **Keep events in subfolders** — KinderSort only reads photos inside named subfolders of the Events folder

---

*KinderSort — Student Photo Organiser | Runs fully offline, no internet required*
"""

    guidebook_path = Path("guidebook.md")
    guidebook_path.write_text(content, encoding="utf-8")
    print(f"guidebook.md written ({len(content)} chars)")


def export_to_docx_python() -> None:
    """Export guidebook.md to KinderSort_Teacher_Guide.docx using python-docx.

    Parses the markdown manually (headings, paragraphs, code blocks, images)
    and builds a Word document with inline images from guidebook_assets/.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("python-docx not installed. Run: pip install python-docx")
        return

    print("\nExporting to .docx...")
    doc = Document()

    # Style the document
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    content = Path("guidebook.md").read_text(encoding="utf-8")
    lines = content.split("\n")

    in_code_block = False
    code_lines = []

    for line in lines:
        # Code block toggle
        if line.startswith("```"):
            if in_code_block:
                # End code block — add as a paragraph with monospace font
                para = doc.add_paragraph()
                run = para.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                para.paragraph_format.left_indent = Inches(0.5)
                code_lines = []
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Headings
        if line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        # Inline image
        elif line.startswith("!["):
            # Extract path: ![alt](path)
            try:
                path_part = line.split("(")[1].rstrip(")")
                img_path = Path(path_part)
                if img_path.exists():
                    para = doc.add_paragraph()
                    run = para.add_run()
                    run.add_picture(str(img_path), width=Inches(5.5))
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    doc.add_paragraph(f"[Image not found: {path_part}]")
            except (IndexError, Exception) as e:  # noqa: BLE001
                doc.add_paragraph(f"[Image error: {e}]")
        # Horizontal rule
        elif line.strip() == "---":
            doc.add_paragraph("─" * 60)
        # List items
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("  - ") or line.startswith("  * "):
            doc.add_paragraph(line[4:], style="List Bullet 2")
        # Numbered list
        elif len(line) > 2 and line[0].isdigit() and line[1] == ".":
            doc.add_paragraph(line[3:], style="List Number")
        # Blockquote
        elif line.startswith("> "):
            para = doc.add_paragraph(line[2:])
            para.paragraph_format.left_indent = Inches(0.4)
        # Empty line
        elif not line.strip():
            pass  # Skip blank lines (headings/paragraphs handle spacing)
        # Normal paragraph
        else:
            doc.add_paragraph(line)

    out_path = Path("KinderSort_Teacher_Guide.docx")
    doc.save(str(out_path))
    print(f"Exported: {out_path.resolve()}")


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="KinderSort guidebook generator")
    parser.add_argument(
        "--mode",
        choices=["capture", "write-guide", "export-docx", "all"],
        default="all",
        help="What to run: capture screenshots, write guide, export docx, or all",
    )
    args = parser.parse_args()

    if args.mode in ("capture", "all"):
        run_guide_capture()

    if args.mode in ("write-guide", "all"):
        write_guidebook_md()

    if args.mode in ("export-docx", "all"):
        export_to_docx_python()
